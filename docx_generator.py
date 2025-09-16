import os
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from google_disk import GoogleDriveManager
from typing import Optional, Dict, Any
import logging
import re
#
# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_styled_document():
    """Создает документ Word с корпоративными стилями"""
    doc = Document()
    
    # Настройка стилей документа
    styles = doc.styles
    
    # Основные цвета (соответствуют PDF версии)
    PRIMARY_BLUE = RGBColor(31, 78, 121)  # #1F4E79
    BLACK = RGBColor(0, 0, 0)
    DARK_GRAY = RGBColor(51, 51, 51)  # #333333
    GRAY = RGBColor(85, 85, 85)  # #555555
    
    # Создаем стиль для заголовка документа
    try:
        # Проверяем, существует ли уже стиль
        title_style = styles['CustomTitle']
    except KeyError:
        try:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.base_style = styles['Normal']
            
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(16)
            title_font.bold = True
            title_font.color.rgb = DARK_GRAY
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(20)
        except Exception:
            # Если не удается создать стиль, продолжаем без него
            pass
    
    # Создаем стиль для заголовков секций
    try:
        section_style = styles['SectionHeader']
    except KeyError:
        try:
            section_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
            section_style.base_style = styles['Normal']
            
            section_font = section_style.font
            section_font.name = 'Calibri'
            section_font.size = Pt(14)
            section_font.bold = True
            section_font.color.rgb = PRIMARY_BLUE
            section_style.paragraph_format.space_before = Pt(20)
            section_style.paragraph_format.space_after = Pt(12)
        except Exception:
            pass
    
    # Создаем стиль для подзаголовков
    try:
        subheader_style = styles['SubHeader']
    except KeyError:
        try:
            subheader_style = styles.add_style('SubHeader', WD_STYLE_TYPE.PARAGRAPH)
            subheader_style.base_style = styles['Normal']
            
            subheader_font = subheader_style.font
            subheader_font.name = 'Calibri'
            subheader_font.size = Pt(12)
            subheader_font.bold = True
            subheader_font.color.rgb = BLACK
            subheader_style.paragraph_format.space_before = Pt(8)
            subheader_style.paragraph_format.space_after = Pt(6)
        except Exception:
            pass
    
    # Создаем стиль для основного текста
    try:
        body_style = styles['CustomBody']
    except KeyError:
        try:
            body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
            body_style.base_style = styles['Normal']
            
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
            body_font.color.rgb = BLACK
            body_style.paragraph_format.space_after = Pt(8)
        except Exception:
            pass
    
    # Создаем стиль для вторичного текста
    try:
        secondary_style = styles['SecondaryText']
    except KeyError:
        try:
            secondary_style = styles.add_style('SecondaryText', WD_STYLE_TYPE.PARAGRAPH)
            secondary_style.base_style = styles['Normal']
            
            secondary_font = secondary_style.font
            secondary_font.name = 'Calibri'
            secondary_font.size = Pt(10)
            secondary_font.color.rgb = GRAY
            secondary_style.paragraph_format.space_after = Pt(6)
        except Exception:
            pass
    
    return doc

def process_styled_text_to_docx(doc: Document, text: str, title: str = "Документ"):
    """Обрабатывает текст с HTML-тегами стилизации и добавляет в документ Word"""
    
    # Добавляем заголовок документа
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    
    # Применяем стиль заголовка, если он существует
    try:
        title_paragraph.style = 'CustomTitle'
    except KeyError:
        # Если стиль не найден, применяем форматирование напрямую
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(51, 51, 51)  # DARK_GRAY
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Разбиваем текст на параграфы
    paragraphs = text.split('\n')
    
    for paragraph_text in paragraphs:
        if paragraph_text.strip():
            safe_paragraph = paragraph_text.strip()
            
            # Обрабатываем <br> теги
            if safe_paragraph == '<br>' or safe_paragraph == '<br/>':
                doc.add_paragraph()  # Добавляем пустой параграф
                continue
            
            # Удаляем проблемные символы
            safe_paragraph = safe_paragraph.replace('■', '')
            safe_paragraph = safe_paragraph.replace('\ufffd', '')
            safe_paragraph = safe_paragraph.replace('\u25a0', '')
            safe_paragraph = ' '.join(safe_paragraph.split())
            
            # Удаляем одиночные <br> теги
            safe_paragraph = re.sub(r'<br\s*/?>', '', safe_paragraph, flags=re.IGNORECASE)
            
            if not safe_paragraph:
                continue
            
            # Определяем стиль и обрабатываем HTML теги
            paragraph = doc.add_paragraph()
            
            # Проверяем, является ли это заголовком секции (поддерживаем оба цвета)
            if re.search(r'<b color="#(1F4E79|4A90E2)">', safe_paragraph, re.IGNORECASE):
                try:
                    paragraph.style = 'SectionHeader'
                except KeyError:
                    # Применяем форматирование напрямую
                    pass
                # Убираем HTML теги и делаем текст заглавными буквами
                clean_text = re.sub(r'<b color="#(1F4E79|4A90E2)">(.*?)</b>', r'\2', safe_paragraph, flags=re.IGNORECASE)
                run = paragraph.add_run(clean_text.upper())
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)  # PRIMARY_BLUE
                
            # Проверяем, является ли это подзаголовком
            elif re.search(r'<b>(?!.*color)', safe_paragraph, re.IGNORECASE):
                try:
                    paragraph.style = 'SubHeader'
                except KeyError:
                    pass
                clean_text = re.sub(r'<b>(.*?)</b>', r'\1', safe_paragraph, flags=re.IGNORECASE)
                run = paragraph.add_run(clean_text)
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # BLACK
                
            # Проверяем, содержит ли вторичный текст
            elif re.search(r'color="#555555"', safe_paragraph, re.IGNORECASE):
                try:
                    paragraph.style = 'SecondaryText'
                except KeyError:
                    pass
                # Обрабатываем смешанный контент с тегами
                process_mixed_content(paragraph, safe_paragraph)
                
            else:
                # Обычный текст с возможными цветными вставками
                try:
                    paragraph.style = 'CustomBody'
                except KeyError:
                    pass
                process_mixed_content(paragraph, safe_paragraph)
        else:
            # Добавляем пустой параграф для пустых строк
            doc.add_paragraph()

def process_mixed_content(paragraph, text):
    """Обрабатывает текст со смешанным содержимым (обычный текст + HTML теги)"""
    PRIMARY_BLUE = RGBColor(31, 78, 121)  # #1F4E79
    GRAY = RGBColor(85, 85, 85)  # #555555
    
    # Разбиваем текст на части с учетом HTML тегов
    parts = re.split(r'(<font color="[^"]*">.*?</font>|<b>.*?</b>)', text, flags=re.IGNORECASE)
    
    for part in parts:
        if not part:
            continue
            
        # Обрабатываем цветной текст
        color_match = re.match(r'<font color="([^"]*)">(.*?)</font>', part, re.IGNORECASE)
        if color_match:
            color_value = color_match.group(1)
            content = color_match.group(2)
            
            run = paragraph.add_run(content)
            # Поддерживаем оба цвета для технологий
            if color_value.lower() in ['#1f4e79', '#4a90e2']:
                run.font.color.rgb = PRIMARY_BLUE
            elif color_value.lower() == '#555555':
                run.font.color.rgb = GRAY
            continue
        
        # Обрабатываем жирный текст
        bold_match = re.match(r'<b>(.*?)</b>', part, re.IGNORECASE)
        if bold_match:
            content = bold_match.group(1)
            run = paragraph.add_run(content)
            run.font.bold = True
            continue
        
        # Обычный текст
        if part.strip():
            paragraph.add_run(part)

def create_docx_from_text(text: str, output_path: str, title: str = "Документ") -> bool:
    """
    Создает Word документ из текста с поддержкой стилизации
    
    Args:
        text: Текст для записи в документ (может содержать HTML-теги для стилизации)
        output_path: Путь для сохранения документа
        title: Заголовок документа
    
    Returns:
        True если файл создан успешно, False в случае ошибки
    """
    try:
        # Создаем директорию если она не существует
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Создаем стилизованный документ
        doc = create_styled_document()
        
        # Обрабатываем и добавляем текст
        process_styled_text_to_docx(doc, text, title)
        
        # Сохраняем документ
        doc.save(output_path)
        
        logger.info(f"Word документ успешно создан: {output_path}")
        return True
        
    except Exception as e:
        logger.error(f"Ошибка создания Word документа: {e}")
        return False

def create_docx_bytes_from_text(text: str, title: str = "Документ") -> Optional[bytes]:
    """
    Создает Word документ из текста и возвращает его в виде байтов
    
    Args:
        text: Текст для записи в документ (может содержать HTML-теги для стилизации)
        title: Заголовок документа
    
    Returns:
        Байты Word документа или None в случае ошибки
    """
    try:
        # Создаем стилизованный документ
        doc = create_styled_document()
        
        # Обрабатываем и добавляем текст
        process_styled_text_to_docx(doc, text, title)
        
        # Создаем буфер в памяти
        buffer = io.BytesIO()
        doc.save(buffer)
        
        # Получаем байты
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        logger.info("Word документ успешно создан в памяти")
        return docx_bytes
        
    except Exception as e:
        logger.error(f"Ошибка создания Word документа в памяти: {e}")
        return None

def create_and_upload_docx_to_drive(
    text: str, 
    file_name: str,
    folder_name: Optional[str] = None,
    title: str = "Документ",
    credentials_path: str = "oauth.json"
) -> Dict[str, Any]:
    """
    Создает Word документ из текста и загружает его в Google Drive
    
    Args:
        text: Текст для записи в документ
        file_name: Имя файла (без расширения .docx)
        folder_name: Имя папки в Google Drive (создается если не существует)
        title: Заголовок документа
        credentials_path: Путь к файлу OAuth credentials
    
    Returns:
        Словарь с результатом операции
    """
    try:
        # Добавляем расширение .docx если его нет
        if not file_name.endswith('.docx'):
            file_name += '.docx'
        
        # Создаем Word документ в памяти
        docx_bytes = create_docx_bytes_from_text(text, title)
        if not docx_bytes:
            return {
                'success': False,
                'error': 'Не удалось создать Word документ'
            }
        
        # Инициализируем Google Drive Manager
        drive_manager = GoogleDriveManager(credentials_path=credentials_path)
        
        # Получаем или создаем папку
        folder_id = None
        if folder_name:
            folder_id = drive_manager.get_or_create_folder(folder_name)
            if not folder_id:
                return {
                    'success': False,
                    'error': 'Не удалось создать папку в Google Drive'
                }
        
        # Загружаем Word документ
        upload_result = drive_manager.upload_file_from_bytes(
            file_data=docx_bytes,
            file_name=file_name,
            mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            folder_id=folder_id
        )
        
        if upload_result.get('success'):
            file_id = upload_result.get('file_id')
            
            # Делаем файл общедоступным
            if file_id:
                permissions_set = drive_manager.set_file_permissions(
                    file_id, 
                    permission_type='reader', 
                    role='anyone'
                )
                if permissions_set:
                    logger.info(f"Word документ '{file_name}' успешно загружен и сделан общедоступным")
                else:
                    logger.warning(f"Word документ '{file_name}' загружен, но не удалось сделать его общедоступным")
            
            return upload_result
        else:
            return upload_result
            
    except Exception as e:
        logger.error(f"Ошибка создания и загрузки Word документа: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def save_docx_locally_and_upload(
    text: str,
    file_name: str,
    folder_name: Optional[str] = None,
    title: str = "Документ",
    local_dir: str = "downloads",
    credentials_path: str = "oauth.json"
) -> Dict[str, Any]:
    """
    Создает Word документ локально и загружает его в Google Drive
    
    Args:
        text: Текст для записи в документ
        file_name: Имя файла (без расширения .docx)
        folder_name: Имя папки в Google Drive
        title: Заголовок документа
        local_dir: Локальная директория для сохранения
        credentials_path: Путь к файлу OAuth credentials
    
    Returns:
        Словарь с результатом операции
    """
    try:
        # Добавляем расширение .docx если его нет
        if not file_name.endswith('.docx'):
            file_name += '.docx'
        
        # Создаем локальный путь
        local_path = os.path.join(local_dir, file_name)
        
        # Создаем Word документ локально
        if not create_docx_from_text(text, local_path, title):
            return {
                'success': False,
                'error': 'Не удалось создать Word документ локально'
            }
        
        # Инициализируем Google Drive Manager
        drive_manager = GoogleDriveManager(credentials_path=credentials_path)
        
        # Получаем или создаем папку
        folder_id = None
        if folder_name:
            folder_id = drive_manager.get_or_create_folder(folder_name)
            if not folder_id:
                # Удаляем локальный файл в случае ошибки
                if os.path.exists(local_path):
                    os.remove(local_path)
                return {
                    'success': False,
                    'error': 'Не удалось создать папку в Google Drive'
                }
        
        # Загружаем файл
        upload_result = drive_manager.upload_file(
            file_path=local_path,
            folder_id=folder_id,
            file_name=file_name
        )
        
        if upload_result.get('success'):
            file_id = upload_result.get('file_id')
            
            # Делаем файл общедоступным
            if file_id:
                permissions_set = drive_manager.set_file_permissions(
                    file_id, 
                    permission_type='reader', 
                    role='anyone'
                )
                if permissions_set:
                    logger.info(f"Word документ '{file_name}' успешно загружен и сделан общедоступным")
                else:
                    logger.warning(f"Word документ '{file_name}' загружен, но не удалось сделать его общедоступным")
        
        # Удаляем локальный файл после загрузки
        if os.path.exists(local_path):
            os.remove(local_path)
            logger.info(f"Локальный файл {local_path} удален")
        
        return upload_result
        
    except Exception as e:
        logger.error(f"Ошибка создания и загрузки Word документа: {e}")
        # Удаляем локальный файл в случае ошибки
        local_path = os.path.join(local_dir, file_name)
        if os.path.exists(local_path):
            os.remove(local_path)
        return {
            'success': False,
            'error': str(e)
        }
